"""
Text Extractors
Extract text from various resume document formats (PDF, DOCX)
"""

import PyPDF2
from docx import Document
from pathlib import Path
from typing import Optional
import re
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF using PyPDF2

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If PDF extraction fails
        """
        text = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                logger.info(f"Extracting text from PDF: {file_path}")
                logger.info(f"Total pages: {len(pdf_reader.pages)}")

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                        logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")
                    else:
                        logger.warning(f"No text extracted from page {page_num}")

        except FileNotFoundError:
            raise ValueError(f"PDF file not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Error extracting PDF: {str(e)}")

        combined_text = "\n".join(text)

        if not combined_text.strip():
            logger.warning("PDF extraction resulted in empty text - may be scanned/image-based")
            raise ValueError(
                "No text extracted from PDF. File may be scanned or image-based. "
                "OCR is not supported in MVP."
            )

        logger.info(f"Successfully extracted {len(combined_text)} characters from PDF")
        return combined_text

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX using python-docx

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text as string

        Raises:
            ValueError: If DOCX extraction fails
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")

            doc = Document(file_path)
            text = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            # Extract from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            combined_text = "\n".join(text)

            if not combined_text.strip():
                raise ValueError("No text extracted from DOCX file")

            logger.info(f"Successfully extracted {len(combined_text)} characters from DOCX")
            return combined_text

        except FileNotFoundError:
            raise ValueError(f"DOCX file not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Error extracting DOCX: {str(e)}")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and special characters

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text

        Examples:
            >>> TextExtractor.clean_text("Too   much    space")
            'Too much space'
        """
        if not text:
            return ""

        # Remove null bytes and control characters (except newlines and tabs)
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

        # Replace multiple spaces with single space (but preserve newlines)
        text = re.sub(r' +', ' ', text)

        # Replace multiple newlines with double newline (paragraph separation)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Remove spaces at the beginning and end of lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        # Strip leading/trailing whitespace
        text = text.strip()

        return text

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Universal text extractor - auto-detects format and extracts text

        Args:
            file_path: Path to resume file (PDF or DOCX)

        Returns:
            Cleaned extracted text

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is unsupported or extraction fails

        Examples:
            >>> text = TextExtractor.extract_text("resume.pdf")
            >>> text = TextExtractor.extract_text("resume.docx")
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        # Get file extension
        suffix = path.suffix.lower()
        logger.info(f"Detected file format: {suffix}")

        # Extract based on format
        if suffix == '.pdf':
            raw_text = TextExtractor.extract_from_pdf(file_path)
        elif suffix in ['.docx', '.doc']:
            if suffix == '.doc':
                logger.warning(
                    ".doc format detected. Only .docx is fully supported. "
                    "Please convert to .docx for best results."
                )
            raw_text = TextExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(
                f"Unsupported file format: {suffix}. "
                f"Supported formats: .pdf, .docx"
            )

        # Clean the extracted text
        cleaned_text = TextExtractor.clean_text(raw_text)

        if not cleaned_text:
            raise ValueError("Text extraction resulted in empty content")

        logger.info(f"Final cleaned text length: {len(cleaned_text)} characters")
        return cleaned_text

    @staticmethod
    def get_text_stats(text: str) -> dict:
        """
        Get statistics about extracted text

        Args:
            text: Extracted text

        Returns:
            Dictionary with text statistics

        Examples:
            >>> stats = TextExtractor.get_text_stats(text)
            >>> print(stats['word_count'])
        """
        lines = text.split('\n')
        words = text.split()

        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'avg_words_per_line': len(words) / len(lines) if lines else 0,
            'is_empty': len(text.strip()) == 0
        }


# Convenience functions for quick access
def extract_text(file_path: str) -> str:
    """
    Quick function to extract text from resume file

    Args:
        file_path: Path to PDF or DOCX file

    Returns:
        Cleaned extracted text
    """
    return TextExtractor.extract_text(file_path)


def extract_and_analyze(file_path: str) -> tuple[str, dict]:
    """
    Extract text and return both text and statistics

    Args:
        file_path: Path to PDF or DOCX file

    Returns:
        Tuple of (extracted_text, statistics)
    """
    extractor = TextExtractor()
    text = extractor.extract_text(file_path)
    stats = extractor.get_text_stats(text)
    return text, stats
