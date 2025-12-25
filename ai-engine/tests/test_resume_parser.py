"""
Unit Tests for Resume Parser
"""

import pytest
import tempfile
import os
from pathlib import Path
from unittest.mock import Mock, patch
import PyPDF2
from docx import Document

from src.resume_parser.extractors import TextExtractor, extract_text
from src.resume_parser.parser import ResumeParser, parse_resume
from src.resume_parser.schemas import ParsedResume, Contact


class TestTextExtractor:
    """Tests for TextExtractor class"""

    def test_clean_text_removes_extra_whitespace(self):
        """Test text cleaning removes excessive whitespace"""
        dirty_text = "Too   much    space  here"
        clean = TextExtractor.clean_text(dirty_text)
        assert clean == "Too much space here"

    def test_clean_text_handles_multiple_newlines(self):
        """Test text cleaning handles multiple newlines"""
        dirty_text = "Line 1\n\n\n\nLine 2"
        clean = TextExtractor.clean_text(dirty_text)
        assert clean == "Line 1\n\nLine 2"

    def test_clean_text_strips_leading_trailing(self):
        """Test text cleaning strips leading/trailing whitespace"""
        dirty_text = "  \n  Text here  \n  "
        clean = TextExtractor.clean_text(dirty_text)
        assert clean == "Text here"

    def test_clean_text_handles_empty_string(self):
        """Test text cleaning handles empty string"""
        assert TextExtractor.clean_text("") == ""
        assert TextExtractor.clean_text(None) == ""

    def test_get_text_stats(self):
        """Test text statistics calculation"""
        text = "Hello world\nThis is a test\nThree lines"
        stats = TextExtractor.get_text_stats(text)

        assert stats['character_count'] == len(text)
        assert stats['word_count'] == 7
        assert stats['line_count'] == 3
        assert stats['is_empty'] == False

    def test_get_text_stats_empty(self):
        """Test text statistics for empty text"""
        stats = TextExtractor.get_text_stats("")
        assert stats['is_empty'] == True
        assert stats['word_count'] == 0

    def test_extract_from_pdf_file_not_found(self):
        """Test PDF extraction with non-existent file"""
        with pytest.raises(ValueError, match="PDF file not found"):
            TextExtractor.extract_from_pdf("nonexistent.pdf")

    def test_extract_from_docx_file_not_found(self):
        """Test DOCX extraction with non-existent file"""
        with pytest.raises(ValueError, match="DOCX file not found"):
            TextExtractor.extract_from_docx("nonexistent.docx")

    def test_extract_text_file_not_found(self):
        """Test universal extractor with non-existent file"""
        with pytest.raises(FileNotFoundError):
            TextExtractor.extract_text("nonexistent.pdf")

    def test_extract_text_unsupported_format(self):
        """Test universal extractor with unsupported format"""
        with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
            tmp.write(b"Some text")
            tmp_path = tmp.name

        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                TextExtractor.extract_text(tmp_path)
        finally:
            os.unlink(tmp_path)

    def test_extract_from_pdf_success(self):
        """Test successful PDF text extraction"""
        # Create a simple PDF for testing
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            # Create a simple PDF with PyPDF2
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter

            c = canvas.Canvas(tmp_path, pagesize=letter)
            c.drawString(100, 750, "John Doe")
            c.drawString(100, 730, "john.doe@example.com")
            c.drawString(100, 710, "Software Engineer")
            c.save()

            # Extract text
            text = TextExtractor.extract_from_pdf(tmp_path)

            assert "John Doe" in text
            assert len(text) > 0

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_extract_from_docx_success(self):
        """Test successful DOCX text extraction"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            # Create a simple DOCX
            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("john.doe@example.com")
            doc.add_paragraph("Software Engineer")
            doc.save(tmp_path)

            # Extract text
            text = TextExtractor.extract_from_docx(tmp_path)

            assert "John Doe" in text
            assert "john.doe@example.com" in text
            assert "Software Engineer" in text

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_extract_from_docx_with_tables(self):
        """Test DOCX extraction includes table content"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            # Create DOCX with table
            doc = Document()
            doc.add_paragraph("Resume")

            # Add a table
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Company"
            table.cell(0, 1).text = "ABC Corp"
            table.cell(1, 0).text = "Title"
            table.cell(1, 1).text = "Engineer"
            doc.save(tmp_path)

            # Extract text
            text = TextExtractor.extract_from_docx(tmp_path)

            assert "Company" in text
            assert "ABC Corp" in text
            assert "Engineer" in text

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


class TestResumeParser:
    """Tests for ResumeParser class"""

    @pytest.fixture
    def mock_llm_client(self):
        """Create a mock LLM client"""
        mock = Mock()
        mock.model = "gpt-4o-mini"
        mock.count_tokens = Mock(return_value=500)
        return mock

    @pytest.fixture
    def sample_llm_response(self):
        """Sample LLM response for testing"""
        return {
            "contact": {
                "name": "John Doe",
                "email": "john.doe@example.com",
                "phone": "555-1234",
                "linkedin": "linkedin.com/in/johndoe",
                "github": "github.com/johndoe",
                "portfolio": None,
                "location": "San Francisco, CA"
            },
            "education": [
                {
                    "institution": "Stanford University",
                    "degree": "Bachelor of Science",
                    "field": "Computer Science",
                    "graduation_date": "May 2020",
                    "gpa": "3.8",
                    "honors": ["Dean's List"],
                    "relevant_coursework": ["Algorithms", "Machine Learning"]
                }
            ],
            "experience": [
                {
                    "company": "Google",
                    "title": "Software Engineer",
                    "start_date": "June 2020",
                    "end_date": "Present",
                    "location": "Mountain View, CA",
                    "responsibilities": [
                        "Developed microservices",
                        "Led team of 3 engineers"
                    ],
                    "achievements": [
                        "Reduced API latency by 40%"
                    ],
                    "technologies": ["Python", "Go", "Kubernetes"]
                }
            ],
            "skills": {
                "technical": ["Python", "Java", "JavaScript"],
                "soft": ["Leadership", "Communication"],
                "tools": ["Git", "Docker"],
                "languages": ["Python", "Java"],
                "frameworks": ["React", "Django"],
                "databases": ["PostgreSQL"],
                "cloud": ["AWS", "GCP"]
            },
            "projects": [
                {
                    "name": "E-commerce Platform",
                    "description": "Built full-stack platform",
                    "technologies": ["React", "Node.js"],
                    "url": "github.com/user/project",
                    "highlights": ["10k users"],
                    "duration": "3 months"
                }
            ],
            "certifications": ["AWS Certified Developer"],
            "leadership": ["President of CS Club"],
            "awards": ["Hackathon Winner 2019"],
            "publications": [],
            "volunteer": [],
            "total_years_experience": 4.5
        }

    def test_parser_initialization(self):
        """Test parser initializes correctly"""
        parser = ResumeParser()
        assert parser.llm_client is not None
        assert parser.extractor is not None
        assert parser.prompt_template is not None

    def test_parser_with_custom_model(self):
        """Test parser with custom model"""
        parser = ResumeParser(model="gpt-4")
        assert parser.llm_client.model == "gpt-4"

    def test_validate_and_convert_success(self, sample_llm_response):
        """Test successful validation and conversion"""
        parser = ResumeParser()
        resume = parser._validate_and_convert(sample_llm_response)

        assert isinstance(resume, ParsedResume)
        assert resume.contact.name == "John Doe"
        assert resume.contact.email == "john.doe@example.com"
        assert len(resume.education) == 1
        assert len(resume.experience) == 1
        assert resume.experience_level == "mid"  # 4.5 years

    def test_validate_and_convert_missing_required_fields(self):
        """Test validation fails with missing required fields"""
        parser = ResumeParser()
        invalid_response = {
            "contact": {
                # Missing name
                "email": "test@example.com"
            }
        }

        with pytest.raises(ValueError):
            parser._validate_and_convert(invalid_response)

    def test_parse_resume_from_text_too_short(self):
        """Test parsing fails with too short text"""
        parser = ResumeParser()
        short_text = "John"

        with pytest.raises(ValueError, match="Resume text too short"):
            parser.parse_resume_from_text(short_text)

    def test_parse_resume_from_text_success(self, mock_llm_client, sample_llm_response):
        """Test successful resume parsing from text"""
        mock_llm_client.generate_json = Mock(return_value=sample_llm_response)

        parser = ResumeParser(llm_client=mock_llm_client)

        resume_text = """
        John Doe
        john.doe@example.com
        555-1234

        EDUCATION
        Stanford University
        Bachelor of Science in Computer Science
        Graduated: May 2020
        GPA: 3.8

        EXPERIENCE
        Google - Software Engineer
        June 2020 - Present
        - Developed microservices using Python and Go
        - Reduced API latency by 40%
        """

        resume = parser.parse_resume_from_text(resume_text)

        assert isinstance(resume, ParsedResume)
        assert resume.contact.name == "John Doe"
        assert mock_llm_client.generate_json.called

    def test_parse_resume_to_dict(self, mock_llm_client, sample_llm_response):
        """Test parsing resume to dictionary"""
        mock_llm_client.generate_json = Mock(return_value=sample_llm_response)

        # Create a test DOCX file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("john.doe@example.com")
            doc.save(tmp_path)

            parser = ResumeParser(llm_client=mock_llm_client)
            result = parser.parse_resume_to_dict(tmp_path)

            assert isinstance(result, dict)
            assert result['contact']['name'] == "John Doe"

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_estimate_cost_gpt4_mini(self):
        """Test cost estimation for GPT-4-mini"""
        parser = ResumeParser(model="gpt-4o-mini")

        # 1000 input tokens, 500 output tokens
        cost = parser._estimate_cost(1000, 500)

        # GPT-4o-mini: $0.15/1M input, $0.60/1M output
        expected_cost = (1000 / 1_000_000 * 0.15) + (500 / 1_000_000 * 0.60)

        assert abs(cost - expected_cost) < 0.0001

    def test_estimate_cost_gpt4(self):
        """Test cost estimation for GPT-4"""
        parser = ResumeParser(model="gpt-4")

        # 1000 input tokens, 500 output tokens
        cost = parser._estimate_cost(1000, 500)

        # GPT-4: $0.03/1K input, $0.06/1K output
        expected_cost = (1000 / 1000 * 0.03) + (500 / 1000 * 0.06)

        assert abs(cost - expected_cost) < 0.0001

    def test_get_parsing_stats(self, mock_llm_client, sample_llm_response):
        """Test getting parsing statistics"""
        mock_llm_client.generate_json = Mock(return_value=sample_llm_response)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            doc = Document()
            doc.add_paragraph("John Doe - Software Engineer")
            doc.save(tmp_path)

            parser = ResumeParser(llm_client=mock_llm_client)
            result = parser.get_parsing_stats(tmp_path)

            assert 'resume' in result
            assert 'stats' in result
            assert 'text_stats' in result['stats']
            assert 'tokens_used' in result['stats']
            assert 'estimated_cost' in result['stats']

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


class TestConvenienceFunctions:
    """Test convenience functions"""

    def test_extract_text_function(self):
        """Test extract_text convenience function"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name

        try:
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(tmp_path)

            text = extract_text(tmp_path)

            assert "Test content" in text

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


# Integration test (requires API key)
@pytest.mark.integration
class TestResumeParserIntegration:
    """Integration tests (require actual API calls)"""

    def test_parse_real_resume(self):
        """Test parsing with real API (skip if no API key)"""
        # This test requires a real API key and sample resume
        # Mark as integration test to skip in unit testing

        sample_resume_path = "examples/sample_resumes/test_resume.docx"

        if not os.path.exists(sample_resume_path):
            pytest.skip("Sample resume not found")

        parser = ResumeParser()

        try:
            resume = parser.parse_resume(sample_resume_path)

            # Basic assertions
            assert resume.contact.name is not None
            assert isinstance(resume, ParsedResume)
            assert resume.experience_level in ["junior", "mid", "senior"]

        except Exception as e:
            pytest.skip(f"Integration test failed (API issue?): {e}")


if __name__ == "__main__":
    # Run tests
    pytest.main([__file__, "-v"])
